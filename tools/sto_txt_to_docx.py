#!/usr/bin/env python3
"""txt бланков СТО → DOCX по docs/sto-templates/DOC-STYLE.md (кегль ≥ 11pt, таблицы)."""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips, Cm

ROOT = Path(__file__).resolve().parents[1]
TXT_DIR = ROOT / "api/assets/sto-templates/txt"
DOCX_DIR = ROOT / "api/assets/sto-templates/docx"
RAW_DIR = ROOT / "docs/sto-templates/raw"

MM = 56.7


def set_run_font(run, *, size_pt: float = 11, bold: bool = False, italic: bool = False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def add_para(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    space_after: float = 4,
    space_before: float = 0,
):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, italic=italic)
    return p


def set_cell(cell, text: str, *, bold: bool = False, center: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text or ""))
    set_run_font(run, size_pt=11, bold=bold)
    # чуть плотнее в ячейках
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)


def style_table(table):
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "111111")
        borders.append(el)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + max(1, len(rows)), cols=len(headers))
    style_table(table)
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, center=True)
    data = rows if rows else [[""] * len(headers)]
    # гарантируем число строк
    while len(table.rows) < 1 + len(data):
        table.add_row()
    for i, row in enumerate(data):
        for j in range(len(headers)):
            val = row[j] if j < len(row) else ""
            set_cell(table.rows[i + 1].cells[j], val, center=(j == 0 and str(val).isdigit()))
    # пустой абзац после таблицы
    add_para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    return table


def is_section_start(s: str) -> bool:
    return bool(re.match(r"^\d+\.\d*\s", s) or re.match(r"^\d+\.\s+\S", s))


def try_consume_table(lines: list[str], i: int) -> tuple[list[str], list[list[str]], int] | None:
    """Если с позиции i начинается таблица — вернуть (headers, rows, next_i)."""
    s = lines[i].strip()
    n = len(lines)

    # —— Приказ: перечень форм ——
    if s == "№ п/п" and i + 2 < n:
        h2, h3 = lines[i + 1].strip(), lines[i + 2].strip()
        if "Наименование" in h2 and "Приложение" in h3:
            headers = [s, h2, h3]
            j = i + 3
            rows: list[list[str]] = []
            while j < n:
                a = lines[j].strip()
                if not a:
                    j += 1
                    continue
                if is_section_start(a) and not re.fullmatch(r"\d+", a):
                    break
                if re.fullmatch(r"\d+", a) and j + 2 < n:
                    rows.append([a, lines[j + 1].strip(), lines[j + 2].strip()])
                    j += 3
                    continue
                break
            return headers, rows, j

    # —— Регламент: проверка контрагента ——
    if s == "Что проверяем" and i + 2 < n:
        h2, h3 = lines[i + 1].strip(), lines[i + 2].strip()
        if h2.startswith("Где") and "приобщаем" in h3:
            headers = [s, h2, h3]
            j = i + 3
            rows = []
            while j + 2 < n:
                a, b, c = lines[j].strip(), lines[j + 1].strip(), lines[j + 2].strip()
                if not a:
                    j += 1
                    break
                if is_section_start(a):
                    break
                rows.append([a, b, c])
                j += 3
            return headers, rows, j

    # —— Лист ознакомления ——
    if s == "№" and i + 3 < n:
        # Должность / Ф. И. О. / Дата… / Подпись
        block = [lines[k].strip() for k in range(i, min(i + 5, n))]
        if block[1] == "Должность" and block[2].startswith("Ф."):
            headers = block[:5] if len(block) >= 5 and ("Подпись" in block[4] or "Дата" in block[3]) else block[:4]
            if len(headers) < 4:
                return None
            # добрать 5-ю колонку если есть
            if len(headers) == 4 and i + 4 < n and "Подпись" in lines[i + 4]:
                headers = block[:5]
            j = i + len(headers)
            rows = []
            # строки: номер, опционально должность, пустые поля
            while j < n:
                a = lines[j].strip()
                if not a:
                    j += 1
                    # несколько пустых — следующая строка листа
                    continue
                if a.startswith("Регламент подготовлен") or a.startswith("{{") or a == "Утверждаю":
                    break
                if is_section_start(a) and not re.fullmatch(r"\d+", a):
                    break
                if re.fullmatch(r"\d+", a):
                    role = ""
                    k = j + 1
                    while k < n and not lines[k].strip():
                        k += 1
                    if k < n and lines[k].strip() and not re.fullmatch(r"\d+", lines[k].strip()):
                        nxt = lines[k].strip()
                        if not nxt.startswith("Регламент") and not is_section_start(nxt):
                            role = nxt
                            j = k + 1
                        else:
                            j = k
                    else:
                        j = k
                    row = [a, role] + [""] * (len(headers) - 2)
                    rows.append(row[: len(headers)])
                    continue
                break
            if not rows:
                rows = [[str(x), ""] + [""] * (len(headers) - 2) for x in range(1, 5)]
            return headers, rows, j

    # —— Таблицы работ / ЗЧ: № + Наименование… ——
    if s == "№" and i + 1 < n and lines[i + 1].strip().startswith("Наименование"):
        headers = []
        j = i
        while j < n:
            t = lines[j].strip()
            if not t:
                j += 1
                break
            # продолжение заголовка в скобках на след. строке
            if t.startswith("(") and headers:
                headers[-1] = (headers[-1] + " " + t).strip()
                j += 1
                continue
            if re.fullmatch(r"\d+", t) or t.startswith("{{") or t.startswith("Итого"):
                break
            if is_section_start(t):
                break
            headers.append(t)
            j += 1
            if len(headers) >= 8:
                break
        if len(headers) < 2:
            return None
        rows = []
        # пустые строки под заполнение + макросы
        empty_budget = 3
        while j < n:
            t = lines[j].strip()
            if t.startswith("Итого") or is_section_start(t) or t.startswith("Состояние"):
                break
            if not t:
                j += 1
                continue
            if t.startswith("{{Таблица") or t.startswith("{{Итого"):
                # макрос в первой ячейке первой строки
                row = [t] + [""] * (len(headers) - 1)
                rows.append(row)
                j += 1
                continue
            if re.fullmatch(r"\d+", t):
                # номер строки + возможно пустые ячейки до следующего номера / Итого
                row = [t] + [""] * (len(headers) - 1)
                rows.append(row)
                j += 1
                # пропуск пустых под строку
                while j < n and not lines[j].strip():
                    j += 1
                continue
            j += 1
        while len(rows) < empty_budget:
            rows.append([str(len(rows) + 1)] + [""] * (len(headers) - 1))
        return headers, rows, j

    return None


def classify(line: str, idx: int, lines: list[str]) -> str:
    s = line.strip()
    if not s:
        return "blank"
    if s in ("Регламент", "Приказываю:") or s.startswith("Договор"):
        return "title"
    if re.match(r"^Приказ\s*№", s, re.I):
        return "title"
    if re.match(r"^Утверждаю$", s, re.I):
        return "h1"
    if re.match(r"^Лист ознакомления", s, re.I):
        return "title"
    if re.match(r"^\d+\.\s+\S", s) and not re.match(r"^\d+\.\d+", s):
        return "h1"
    if re.match(r"^\d+\.\d+", s):
        return "body"
    # шапка приказа / подзаголовок регламента — по центру только заголовочные строки
    if idx < 8 and (
        "работы мастера" in s.lower()
        or "по заключению" in s.lower()
        or "об утверждении" in s.lower()
        or "и регламента" in s.lower()
        or "и ремонту" in s.lower()
        or "форм договоров" in s.lower()
        or s.startswith("г. ")
        or s.startswith("{{ДатаПриказа}}")
        or (s.startswith("«") and "20" in s)
    ):
        return "center"
    # реквизиты и подпись внизу — обычный абзац слева
    if (
        s.startswith("{{Организация")
        or s.startswith("Индивидуальный предприниматель")
        or s.startswith("ОГРНИП")
        or s.startswith("Адрес")
        or s.startswith("Тел.")
        or s.startswith("E-mail")
        or s.startswith("Режим работы")
        or s.startswith("_______________")
        or "ОрганизацияКратко" in s
    ):
        return "body"
    if s.startswith("Экземпляр") or (s.startswith("(") and "подпись" in s.lower()):
        return "muted"
    return "body"


def build_docx(text: str, out: Path) -> None:
    text = text.replace("\u2028", "\n").replace("\r\n", "\n")
    lines = text.split("\n")
    doc = Document()
    for section in doc.sections:
        section.page_width = Twips(11906)
        section.page_height = Twips(16838)
        section.left_margin = Twips(int(20 * MM))
        section.right_margin = Twips(int(15 * MM))
        section.top_margin = Twips(int(15 * MM))
        section.bottom_margin = Twips(int(15 * MM))

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    i = 0
    while i < len(lines):
        s = lines[i].strip()

        # таблицы
        parsed = try_consume_table(lines, i)
        if parsed:
            headers, rows, nxt = parsed
            add_table(doc, headers, rows)
            i = nxt
            continue

        # город + дата
        if s.startswith("г. ") and i + 1 < len(lines):
            nxt = lines[i + 1].strip()
            if nxt.startswith("«") or nxt.startswith("{{Дата") or re.search(r"\d{4}\s*г", nxt):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(8)
                run1 = p.add_run(s)
                set_run_font(run1, size_pt=11)
                run_tab = p.add_run("\t")
                set_run_font(run_tab, size_pt=11)
                pPr = p._p.get_or_add_pPr()
                tabs = OxmlElement("w:tabs")
                tab = OxmlElement("w:tab")
                tab.set(qn("w:val"), "right")
                tab.set(qn("w:pos"), str(int(170 * MM)))
                tabs.append(tab)
                pPr.append(tabs)
                run2 = p.add_run(nxt)
                set_run_font(run2, size_pt=11)
                i += 2
                continue

        kind = classify(lines[i], i, lines)
        if kind == "blank":
            add_para(doc, "", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
        elif kind == "title":
            add_para(
                doc,
                s,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                size=14 if s in ("Регламент",) or s.startswith("Договор") else 13,
                bold=True,
                space_after=6,
                space_before=4,
            )
        elif kind == "center_bold":
            add_para(doc, s, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, space_after=4)
        elif kind == "center":
            add_para(doc, s, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=False, space_after=2)
        elif kind == "h1":
            add_para(doc, s, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, bold=True, space_before=10, space_after=4)
        elif kind == "muted":
            add_para(doc, s, align=WD_ALIGN_PARAGRAPH.LEFT, size=11, italic=True, space_after=2)
        else:
            add_para(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11, bold=False, space_after=4)
        i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--all", action="store_true")
    ap.add_argument("txt", nargs="*")
    args = ap.parse_args()
    files: list[Path] = []
    if args.all:
        files = sorted(TXT_DIR.glob("*.txt"))
    else:
        for t in args.txt:
            p = Path(t)
            if not p.is_absolute():
                p = TXT_DIR / t
            files.append(p)
    if not files:
        print("Usage: sto_txt_to_docx.py --all | file.txt …", file=sys.stderr)
        sys.exit(1)
    for src in files:
        if not src.is_file():
            print("skip missing", src)
            continue
        name = src.name.replace(".txt", ".docx")
        out = DOCX_DIR / name
        build_docx(src.read_text(encoding="utf-8"), out)
        if RAW_DIR.is_dir():
            build_docx(src.read_text(encoding="utf-8"), RAW_DIR / name)
        # verify tables
        d = Document(str(out))
        print(f"OK {out.name} tables={len(d.tables)}")


if __name__ == "__main__":
    main()
